import os
import re
import uuid
import time
import random
from typing import Generator, List, Tuple, Optional
from openai import OpenAI
from supabase import create_client
from dotenv import load_dotenv
import psycopg2
from psycopg2 import sql, extras
from pgvector.psycopg2 import register_vector
from pgvector import Vector

# optional imports for parsing
try:
    from docx import Document as DocxDocument
except Exception:
    DocxDocument = None

try:
    import fitz
except Exception:
    fitz = None

import unicodedata

# ---------- improved cleaner ----------
def clean_text_for_db(s: Optional[str], max_len: int = 10000) -> Optional[str]:
    if s is None:
        return None
    s = s.replace("\x00", "")
    s = unicodedata.normalize("NFC", s)
    cleaned_chars = []
    for ch in s:
        if ch in ("\t", "\n", "\r"):
            cleaned_chars.append(ch)
            continue
        cat = unicodedata.category(ch)
        if cat.startswith("C"):
            continue
        cleaned_chars.append(ch)
    out = "".join(cleaned_chars)
    out = re.sub(r"[ \t\f\v]{2,}", " ", out)
    out = out.strip()
    if max_len and len(out) > max_len:
        out = out[: max_len - 14] + " ... [truncated]"
    return out

# --- CONFIG ---
TABLE_NAME = "documents"
VECTOR_DIM = 1536

EMBEDDING_BATCH_SIZE = 16
DB_INSERT_BATCH_SIZE = 64
EMBEDDING_RETRY_BASE = 1.0
EMBEDDING_MAX_RETRIES = 5
DB_RETRY_BASE = 0.5
DB_MAX_RETRIES = 4

load_dotenv()
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
SUPABASE_URL = os.getenv("SUPABASE_URL")
SUPABASE_KEY = os.getenv("SUPABASE_KEY")
DB_URL = os.getenv("DB_URL")
BUCKET_NAME = "documents"

client = OpenAI(api_key=OPENAI_API_KEY) if OPENAI_API_KEY else None
supabase = create_client(SUPABASE_URL, SUPABASE_KEY) if SUPABASE_URL and SUPABASE_KEY else None


# ------------------ Postgres / pgvector helpers ------------------
def ensure_pgvector_extension(conn) -> None:
    """Create pgvector extension in public schema (idempotent). Commit afterwards."""
    cur = conn.cursor()
    try:
        cur.execute("CREATE EXTENSION IF NOT EXISTS vector WITH SCHEMA public;")
        conn.commit()
    except Exception as e:
        # If create extension fails (e.g. managed DB without permission), raise a helpful error
        conn.rollback()
        raise RuntimeError(
            "Could not create pgvector extension on the database. Ensure pgvector is installed on the server and your DB user has CREATE EXTENSION privileges. Original error: "
            + str(e)
        )
    finally:
        cur.close()


def ensure_table_exists():
    """Create documents table with embedding vector column.

    Important: create extension first, commit, then register_vector on the client side.
    """
    if not DB_URL:
        raise RuntimeError("DB_URL not configured. Set DB_URL environment variable.")

    conn = psycopg2.connect(DB_URL)
    try:
        # 1) ensure extension is installed on DB (this must run in DB where app connects)
        try:
            ensure_pgvector_extension(conn)
        except Exception as e:
            # bubble helpful message
            print("Error creating extension:", e)
            raise

        # 2) register client-side adapter so psycopg2 knows how to adapt Vector
        try:
            register_vector(conn)
        except Exception as e:
            # register_vector should normally work now that extension exists
            print("Warning: register_vector failed after creating extension:", e)

        cur = conn.cursor()
        create_table_sql = sql.SQL(
            f"""
            CREATE TABLE IF NOT EXISTS {TABLE_NAME} (
                id SERIAL PRIMARY KEY,
                filename TEXT,
                url TEXT,
                chunk_text TEXT,
                embedding VECTOR({VECTOR_DIM}),
                created_at TIMESTAMPTZ DEFAULT now()
            );
            """
        )
        cur.execute(create_table_sql)
        conn.commit()
        cur.close()
        print(f"Table '{TABLE_NAME}' is ready (includes chunk_text).")
    finally:
        conn.close()


# ------------------ Supabase helpers (unchanged, defensive) ------------------
def create_bucket_if_not_exists(bucket_name=BUCKET_NAME):
    if supabase is None:
        print("Supabase not configured; skipping bucket creation.")
        return
    try:
        try:
            info = supabase.storage.get_bucket(bucket_name)
            if info:
                print(f"Bucket '{bucket_name}' already exists or check succeeded.")
                return
        except Exception:
            pass
        resp = supabase.storage.create_bucket(bucket_name, options={"public": True})
        print(f"Bucket create response: {resp}")
    except Exception as e:
        print(f"Warning: could not ensure bucket exists: {e}")


# ------------------ Embeddings ------------------

def generate_embedding_batch(texts: List[str], model="text-embedding-3-small") -> List[List[float]]:
    if client is None:
        raise RuntimeError("OpenAI client not configured (OPENAI_API_KEY missing).")
    attempt = 0
    while True:
        attempt += 1
        try:
            resp = client.embeddings.create(model=model, input=texts)
            embeddings = [r.embedding for r in resp.data]
            if len(embeddings) != len(texts):
                raise RuntimeError("Returned embeddings count mismatch")
            return embeddings
        except Exception as e:
            if attempt >= EMBEDDING_MAX_RETRIES:
                raise RuntimeError(f"Embedding generation failed after retries: {e}")
            backoff = EMBEDDING_RETRY_BASE * (2 ** (attempt - 1)) + random.random() * 0.5
            print(f"Embedding batch attempt {attempt} failed: {e}. Backing off {backoff:.1f}s")
            time.sleep(backoff)


# ------------------ Chunk streaming (unchanged) ------------------
def _words_from_text_stream(text: str):
    if not text:
        return
    for w in re.split(r"\s+", text.strip()):
        if w:
            yield w


def chunk_generator_from_file(file_path: str, chunk_size: int = 500, chunk_overlap: int = 50) -> Generator[str, None, None]:
    if not os.path.isfile(file_path):
        raise FileNotFoundError(f"No such file: {file_path}")
    ext = os.path.splitext(file_path)[1].lower().lstrip(".")
    if ext not in ("pdf", "docx"):
        raise ValueError("Only .pdf and .docx files are supported by this function.")
    carry: List[str] = []
    def _maybe_yield_from_carry():
        nonlocal carry
        while len(carry) >= chunk_size:
            chunk_words = carry[:chunk_size]
            yield " ".join(chunk_words).strip()
            carry = carry[chunk_size - chunk_overlap :]
    if ext == "pdf":
        if fitz is None:
            raise RuntimeError("PyMuPDF (fitz) is required for PDF parsing. Install with `pip install PyMuPDF`.")
        doc = fitz.open(file_path)
        try:
            for page in doc:
                page_text = page.get_text("text") or ""
                for line in page_text.splitlines():
                    for w in _words_from_text_stream(line):
                        carry.append(w)
                        if len(carry) >= chunk_size:
                            for out in _maybe_yield_from_carry():
                                yield out
            if carry:
                yield " ".join(carry).strip()
        finally:
            doc.close()
    else:
        if DocxDocument is None:
            raise RuntimeError("python-docx is required for DOCX parsing. Install with `pip install python-docx`.")
        doc = DocxDocument(file_path)
        for para in doc.paragraphs:
            text = para.text or ""
            for line in text.splitlines():
                for w in _words_from_text_stream(line):
                    carry.append(w)
                    if len(carry) >= chunk_size:
                        for out in _maybe_yield_from_carry():
                            yield out
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text or ""
                    for line in text.splitlines():
                        for w in _words_from_text_stream(line):
                            carry.append(w)
                            if len(carry) >= chunk_size:
                                for out in _maybe_yield_from_carry():
                                    yield out
        if carry:
            yield " ".join(carry).strip()


# ------------------ DB insert ------------------
def store_metadata_and_embedding_batch(
    rows: List[Tuple[str, str, str, Vector]], table_name: str = TABLE_NAME, db_url: str = DB_URL
):
    if not rows:
        return []
    sanitized_rows = []
    for fn, url, chunk_txt, emb in rows:
        safe_fn = clean_text_for_db(fn) if isinstance(fn, str) else fn
        safe_url = clean_text_for_db(url) if isinstance(url, str) else url
        safe_chunk = clean_text_for_db(chunk_txt) if isinstance(chunk_txt, str) else chunk_txt
        sanitized_rows.append((safe_fn, safe_url, safe_chunk, emb))

    attempt = 0
    inserted_ids = []
    while True:
        attempt += 1
        conn = None
        try:
            conn = psycopg2.connect(db_url)
            # ensure extension exists and register adapter on this connection
            try:
                ensure_pgvector_extension(conn)
            except Exception as e:
                # If extension can't be created here, it's still possible it's present; continue but warn
                print("Warning: ensure_pgvector_extension failed (may already exist or lacks permission):", e)
            try:
                register_vector(conn)
            except Exception as e:
                print("Warning: register_vector failed:", e)

            cur = conn.cursor()
            insert_query = sql.SQL(
                f"INSERT INTO {table_name} (filename, url, chunk_text, embedding) VALUES %s RETURNING id;"
            )
            extras.execute_values(
                cur,
                insert_query.as_string(conn),
                sanitized_rows,
                template="(%s, %s, %s, %s)",
                page_size=DB_INSERT_BATCH_SIZE,
            )
            ids = [r[0] for r in cur.fetchall()]
            conn.commit()
            cur.close()
            conn.close()
            inserted_ids.extend(ids)
            return inserted_ids
        except Exception as e:
            print(f"DB batch insert attempt {attempt} failed: {e}")
            try:
                if conn:
                    conn.rollback()
                    conn.close()
            except Exception:
                pass
            if attempt >= DB_MAX_RETRIES:
                raise RuntimeError(f"DB insert failed after retries: {e}")
            backoff = DB_RETRY_BASE * (2 ** (attempt - 1)) + random.random() * 0.3
            time.sleep(backoff)


# ------------------ ingest flow ------------------
def ingest_file(
    file_path,
    bucket_name=BUCKET_NAME,
    chunk_size=500,
    chunk_overlap=50,
    upload=True,
    overwrite=True,
    embedding_batch_size=EMBEDDING_BATCH_SIZE,
    db_insert_batch_size=DB_INSERT_BATCH_SIZE,
):
    create_bucket_if_not_exists(bucket_name)
    ensure_table_exists()

    ext = os.path.splitext(file_path)[1].lower().lstrip(".")
    if ext not in ("pdf", "docx"):
        raise ValueError("ingest_file only supports .pdf and .docx files")

    file_name = os.path.basename(file_path)
    chunks_iter = chunk_generator_from_file(file_path, chunk_size=chunk_size, chunk_overlap=chunk_overlap)

    pending_chunk_labels: List[str] = []
    pending_chunks_metadata: List[Tuple[str, str]] = []
    all_inserted_ids = []
    total_chunks_processed = 0
    batch_count = 0

    def flush_embedding_and_insert(batch_chunk_labels: List[str], batch_metadata: List[Tuple[str, str]]):
        nonlocal all_inserted_ids, batch_count
        if not batch_chunk_labels:
            return
        batch_count += 1
        print(f"[batch {batch_count}] Requesting embeddings for {len(batch_chunk_labels)} chunks...")
        embeddings = generate_embedding_batch(batch_chunk_labels)
        rows = []
        for (chunk_label, _meta), emb in zip(batch_metadata, embeddings):
            rows.append((file_name, "None", chunk_label, Vector(emb)))
        for i in range(0, len(rows), db_insert_batch_size):
            sub = rows[i : i + db_insert_batch_size]
            inserted_ids = store_metadata_and_embedding_batch(sub, table_name=TABLE_NAME, db_url=DB_URL)
            all_inserted_ids.extend(inserted_ids)
        print(f"[batch {batch_count}] Inserted {len(rows)} rows (ids: {len(all_inserted_ids)} total).")

    for i, chunk in enumerate(chunks_iter):
        total_chunks_processed += 1
        raw_chunk_label = f"[chunk {total_chunks_processed}] " + chunk
        chunk_label = clean_text_for_db(raw_chunk_label)
        pending_chunk_labels.append(chunk_label)
        pending_chunks_metadata.append((chunk_label, ""))
        if len(pending_chunk_labels) >= embedding_batch_size:
            flush_embedding_and_insert(pending_chunk_labels, pending_chunks_metadata)
            pending_chunk_labels = []
            pending_chunks_metadata = []

    if pending_chunk_labels:
        flush_embedding_and_insert(pending_chunk_labels, pending_chunks_metadata)

    print(f"Completed processing. Total chunks processed: {total_chunks_processed}. Inserted rows: {len(all_inserted_ids)}")
    return all_inserted_ids


# ------------------ query ------------------

def query_similar_embeddings(query_text, top_k=5):
    try:
        query_embedding = generate_embedding_batch([query_text])[0]
        conn = psycopg2.connect(DB_URL)
        try:
            # try to register adapter on this connection
            try:
                register_vector(conn)
            except Exception:
                print("Warning: register_vector failed. Ensure pgvector is installed and connection is ok.")

            cur = conn.cursor()
            cur.execute(
                sql.SQL(
                    """
                    SELECT id, filename, url, chunk_text, embedding, embedding <-> %s AS distance
                    FROM {table}
                    ORDER BY distance
                    LIMIT %s;
                    """
                ).format(table=sql.Identifier(TABLE_NAME)),
                (Vector(query_embedding), top_k),
            )
            rows = cur.fetchall()
            cur.close()
            conn.close()

            results = []
            for r in rows:
                _id, filename, url, chunk_text_val, embedding_val, distance = r
                results.append(
                    {
                        "id": _id,
                        "filename": filename,
                        "url": url,
                        "chunk_text": chunk_text_val,
                        "embedding": embedding_val,
                        "distance": float(distance) if distance is not None else None,
                    }
                )
            return results
        finally:
            try:
                conn.close()
            except Exception:
                pass
    except Exception as e:
        print(f"Error querying embeddings: {e}")
        return []


# ------------------ optional test helper ------------------
def test_vector_insert():
    if not DB_URL:
        print("DB_URL not set; can't run test")
        return
    conn = psycopg2.connect(DB_URL)
    try:
        ensure_pgvector_extension(conn)
        register_vector(conn)
        cur = conn.cursor()
        cur.execute("DROP TABLE IF EXISTS test_vector_table;")
        cur.execute("CREATE TABLE test_vector_table (id SERIAL PRIMARY KEY, v VECTOR(3));")
        cur.execute("INSERT INTO test_vector_table (v) VALUES (%s) RETURNING id;", (Vector([0.1,0.2,0.3]),))
        print("Inserted id:", cur.fetchone())
        conn.commit()
        cur.close()
    finally:
        conn.close()


if __name__ == "__main__":
    # quick demo (uncomment to run)
    # test_vector_insert()
    # inserted = ingest_file(
    #     "D:/clients/Canines/backend/training-doc/dc1.docx",
    #     chunk_size=200,
    #     chunk_overlap=30,
    #     embedding_batch_size=EMBEDDING_BATCH_SIZE,
    #     db_insert_batch_size=DB_INSERT_BATCH_SIZE,
    # )
    # print("Inserted chunk ids:", inserted)

    print([l["chunk_text"] for l in query_similar_embeddings("amino qauantity.", top_k=6)])
